from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOC_PATH = r"c:\Users\kafis\OneDrive - Eotvos Lorand Tudomanyegyetem\Asztal\Kafi Thesis\credit-risk-prediction-ml\documentation\ELTE_Thesis_Documentation.docx"


def _set_run_font(run, *, bold=False, size=12):
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.rFonts if rpr.rFonts is not None else OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    if rfonts.getparent() is None:
        rpr.append(rfonts)


def add_heading(document, text, size=14):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, bold=True, size=size)


def add_body(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, bold=False, size=12)


def add_image_placeholder(document, placeholder_text, instruction_text):
    ph = document.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ph.paragraph_format.line_spacing = 1.5
    ph.paragraph_format.space_after = Pt(2)
    run = ph.add_run(f"[IMAGE PLACEHOLDER: {placeholder_text}]")
    _set_run_font(run, bold=True, size=12)

    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = f"Instruction: {instruction_text}"
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cp.paragraph_format.line_spacing = 1.5
    for r in cp.runs:
        _set_run_font(r, bold=False, size=12)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def append_chapter_2(document):
    add_heading(document, "Chapter 2: User Documentation", size=14)

    add_heading(document, "2.1 System Requirements", size=13)
    add_body(
        document,
        "This application is a browser-based decision-support interface for credit risk prediction. End users do not need to interact with Python code directly once the system is deployed, but the host machine must be able to run a Flask backend and machine-learning dependencies. For stable operation in a local or departmental environment, a modern dual-core CPU, at least 8 GB RAM, and 2 GB free storage are recommended. Although prediction for a single applicant record is computationally light, model loading and SHAP explanation generation benefit from additional memory headroom."
    )
    add_body(
        document,
        "Software prerequisites include Python 3.10 or newer, pip package management, and a modern web browser (Chrome, Edge, or Firefox). The backend depends on libraries such as Flask, scikit-learn, XGBoost, SHAP, LIME, pandas, and numpy. The user interface relies on server-side HTML templates and client-side JavaScript to dynamically construct the form from the backend schema. Network access is only required for initial dependency installation and optional dataset download; day-to-day prediction can run locally after setup is complete."
    )
    add_image_placeholder(
        document,
        "Host machine terminal showing Python version, available memory, and successful dependency check",
        "Open a terminal in the project root, run Python version and package checks, then capture the terminal with the successful outputs visible."
    )

    add_heading(document, "2.2 Installation and Setup", size=13)
    add_body(
        document,
        "The installation process follows a reproducible sequence designed for users who may not be machine-learning specialists. Step 1 is repository acquisition. Clone the project and enter the root directory. Step 2 is environment preparation. Create and activate an isolated virtual environment to avoid package conflicts with system-wide Python modules. Step 3 is dependency installation through requirements.txt. Step 4 is data preparation. Execute the data loader and preprocessing modules so that the training and inference artifacts are generated in their expected directories. Step 5 is model training. Run the training module to produce saved models in the model storage folder. Step 6 is application startup. Launch the Flask application and open the provided URL in the browser."
    )
    add_body(
        document,
        "Environment variables are optional but recommended for controlled deployment. The Flask host, port, and debug behavior can be configured through FLASK_HOST, FLASK_PORT, and FLASK_DEBUG. In production-like settings, the secret key should be set explicitly instead of relying on a default value. This setup pattern allows both a quick local demonstration and a more disciplined deployment profile for shared use."
    )
    add_image_placeholder(
        document,
        "Terminal sequence for clone, virtual environment activation, and dependency installation",
        "Run the full setup command sequence from clone to pip install, then capture the terminal showing each command and a successful completion state."
    )
    add_image_placeholder(
        document,
        "Terminal output for data preparation and model training scripts",
        "Execute preprocessing and training commands, then capture the output that confirms saved arrays, preprocessor artifact, and trained model files."
    )
    add_image_placeholder(
        document,
        "Browser window after Flask startup, showing the application landing page",
        "Start the web application and open the local URL, then capture the first fully loaded page with header and form region visible."
    )

    add_heading(document, "2.3 Interface Overview", size=13)
    add_body(
        document,
        "The interface is implemented as a single-page workflow optimized for guided data entry and immediate feedback. The page begins with a project header and a short welcome section, followed by the prediction workspace. The left panel contains a model-selection dropdown and a dynamically generated input form that is fetched from the backend schema endpoint. The right panel is initially hidden and becomes visible only after a successful prediction request. This panel displays the risk decision and a SHAP-based contribution chart that helps users understand the factors behind the output."
    )
    add_body(
        document,
        "The design supports non-technical use by ensuring that field constraints and valid categorical values are obtained from the backend rather than hardcoded in the browser. Consequently, if schema options change in the backend, the form adapts without requiring manual frontend code edits for each field. This architecture minimizes UI drift and reduces user confusion caused by stale forms."
    )
    add_image_placeholder(
        document,
        "Full interface layout with header, model selector, dynamic form area, and result panel",
        "Load the application in a desktop browser and capture the full viewport so both the input form section and result section are visible."
    )

    add_heading(document, "2.4 Main Features and Functionalities", size=13)
    add_body(
        document,
        "The software is organized around practical user-facing modules that correspond to the actual route and interaction design of the application."
    )

    add_heading(document, "2.4.1 Module A: Model Selection", size=12)
    add_body(
        document,
        "User Flow: After page load, the system automatically queries the models endpoint and populates the model dropdown. The user clicks the Select Model field and chooses one of the available trained models, such as random_forest, logistic_regression, or xgboost. Expected Result: the selected model is included in the prediction request and the backend uses that model for scoring and explanation."
    )
    add_image_placeholder(
        document,
        "Opened model dropdown listing available trained model options",
        "Wait for model loading to finish, open the model dropdown, and capture the menu with all available options expanded."
    )

    add_heading(document, "2.4.2 Module B: Dynamic Input Form", size=12)
    add_body(
        document,
        "User Flow: The frontend calls the schema endpoint and renders numerical and categorical fields automatically. The user fills all required fields, including numeric values such as age and credit amount, and categorical values such as employment status or housing type. Expected Result: the form enforces required inputs and valid option lists so only schema-compliant data can be submitted."
    )
    add_image_placeholder(
        document,
        "Dynamic form with mixed numeric inputs and categorical dropdowns",
        "Capture the form after schema load with multiple fields visible, including both a number input and a select input in populated state."
    )

    add_heading(document, "2.4.3 Module C: Prediction Execution", size=12)
    add_body(
        document,
        "User Flow: The user clicks Predict Credit Risk. The button enters a temporary Predicting state while the browser sends a JSON request to the predict endpoint. Expected Result: when validation succeeds, the backend returns status, class prediction, probability, risk level, model identifier, and explanation payload. The interface then displays the result card."
    )
    add_image_placeholder(
        document,
        "Prediction button in loading state during API request",
        "Submit a valid form and quickly capture the button while it displays the Predicting label."
    )

    add_heading(document, "2.4.4 Module D: Decision Result and Confidence", size=12)
    add_body(
        document,
        "User Flow: After receiving a successful response, the user reads the decision banner in the result panel. The banner shows either a low-risk good-credit outcome or a high-risk bad-credit outcome, together with confidence percentage derived from probability. Expected Result: users obtain an immediately interpretable decision summary suitable for first-level screening."
    )
    add_image_placeholder(
        document,
        "Visible prediction result card showing risk label and confidence",
        "Run a prediction and capture the right-side result panel with the full decision text and confidence value."
    )

    add_heading(document, "2.4.5 Module E: Explainability Chart (SHAP)", size=12)
    add_body(
        document,
        "User Flow: In the same result panel, the user inspects the horizontal bar chart that lists top SHAP feature contributions. Positive contributions indicate increased predicted risk, while negative contributions indicate decreased risk. Expected Result: the user can identify the most influential factors and understand the direction of each feature's effect on the final score."
    )
    add_image_placeholder(
        document,
        "SHAP bar chart with top ten feature contributions and color-coded impact direction",
        "Complete a prediction and capture the chart area so feature names, bars, and tooltip legend context are readable."
    )

    add_heading(document, "2.4.6 Module F: Error Handling and Validation Feedback", size=12)
    add_body(
        document,
        "User Flow: If fields are missing, out of range, or incorrectly typed, the backend returns validation errors and the frontend shows an error alert. Expected Result: users receive actionable feedback, correct the form inputs, and resubmit without restarting the page."
    )
    add_image_placeholder(
        document,
        "Validation error message triggered by incomplete or invalid input",
        "Submit the form with at least one required field missing, then capture the displayed error alert/message presented to the user."
    )

    add_heading(document, "2.5 User Roles and Permissions", size=13)
    add_body(
        document,
        "In the current implementation, the system does not define account-based authorization layers such as Guest, User, and Admin. There is no login route, no role database, and no permission middleware in the Flask application. Functionally, this means all users who can access the web endpoint have identical capabilities: view the form, submit prediction requests, and inspect explanations."
    )
    add_body(
        document,
        "For thesis completeness, this section should be interpreted as not applicable in the present version and as a design opportunity for future extension. A future role model could separate read-only demonstration access from privileged operational actions, such as model switching restrictions, prediction logging controls, or administrative model lifecycle operations."
    )
    add_image_placeholder(
        document,
        "Current application behavior without login or role-specific menu",
        "Capture the main page and backend route list evidence showing no authentication screen and no role-dependent navigation elements."
    )

    add_heading(document, "2.6 Troubleshooting", size=13)
    add_body(
        document,
        "Users may encounter setup and runtime issues that are typical for Python web applications with ML dependencies. If the page does not load, confirm that Flask is running and the configured host and port are correct. If form fields do not appear, verify that the schema endpoint is reachable and not blocked by backend startup failure. If prediction fails, check that trained model artifacts and the serialized preprocessor exist in the expected models directory and that input values remain within allowed ranges."
    )
    add_body(
        document,
        "Dependency-related errors are also common, especially missing packages such as python-docx for documentation tooling or model libraries for inference. These should be resolved by reinstalling requirements in the active virtual environment. For incorrect prediction behavior or unexpected explanations, retrain models from a clean preprocessed dataset and restart the application to clear stale in-memory model instances. The health endpoint can be used as a quick operational check before deeper diagnosis."
    )
    add_body(
        document,
        "Practical quick checks for non-technical users are: open the health route in a browser, ensure model names are listed in the model selector, ensure every required input field is filled, and retry prediction with a known valid test sample. For technical support, preserve terminal logs and the exact input scenario so the issue can be reproduced."
    )
    add_image_placeholder(
        document,
        "Terminal and browser evidence for common failure scenarios and fixes",
        "Capture one screenshot each for backend startup error, missing model file error, and successful recovery after corrective action."
    )


def main():
    document = Document(DOC_PATH)
    append_chapter_2(document)
    document.save(DOC_PATH)
    print(DOC_PATH)


if __name__ == "__main__":
    main()