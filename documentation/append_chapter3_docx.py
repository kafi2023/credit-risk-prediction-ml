from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from pathlib import Path


DOC_PATH = r"c:\Users\kafis\OneDrive - Eotvos Lorand Tudomanyegyetem\Asztal\Kafi Thesis\credit-risk-prediction-ml\documentation\ELTE_Thesis_Documentation.docx"
FIGURE_3_1_PATH = Path(__file__).with_name("thesis_class_uml.png")


def _set_run_font(run, *, bold=False, size=12):
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.rFonts if rpr.rFonts is not None else OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    if rfonts.getparent() is None:
        rpr.append(rfonts)


def add_heading(document, text, size=14):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, bold=True, size=size)


def add_body(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, bold=False, size=12)


def add_image_placeholder(document, placeholder_text, instruction_text):
    ph = document.add_paragraph()
    ph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ph.paragraph_format.line_spacing = 1.5
    ph.paragraph_format.space_after = Pt(2)
    run = ph.add_run(f"[IMAGE PLACEHOLDER: {placeholder_text}]")
    _set_run_font(run, bold=True, size=12)

    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = f"Instruction: {instruction_text}"
    cp = cell.paragraphs[0]
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    cp.paragraph_format.line_spacing = 1.5
    for r in cp.runs:
        _set_run_font(r, bold=False, size=12)

    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def add_captioned_image(document, image_path, figure_title, caption_text):
    figure = document.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.space_after = Pt(2)

    run = figure.add_run()
    if image_path.exists():
        run.add_picture(str(image_path), width=Inches(6.3))
    else:
        run.add_text(f"[Missing figure file: {image_path.name}]")

    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.space_after = Pt(6)

    caption_run = caption.add_run(f"Figure 3.1: {figure_title}\n")
    _set_run_font(caption_run, bold=True, size=11)

    body_run = caption.add_run(caption_text)
    _set_run_font(body_run, bold=False, size=10)


def add_pseudocode_block(document, title, lines):
    add_body(document, title)
    for line in lines:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(line)
        _set_run_font(run, bold=False, size=11)
    spacer = document.add_paragraph()
    spacer.paragraph_format.space_after = Pt(6)


def append_chapter_3(document):
    add_heading(document, "Chapter 3: Developer Documentation", size=14)

    add_heading(document, "3.1 System Architecture", size=13)
    add_body(
        document,
        "The implemented system follows a layered pipeline architecture with clear separation between presentation, orchestration, machine-learning workflow, and persistence artifacts. At runtime, the web interface acts as the presentation layer, while Flask routes provide an application service layer that orchestrates request handling. The prediction core resides in the CreditRiskPredictor class, which coordinates input validation, feature transformation through a persisted preprocessor, model inference, and post-hoc explainability. In architectural terms, this combines service-layer orchestration with a modular domain pipeline, rather than a strict MVC pattern, because business logic is concentrated in reusable Python modules outside the route handlers."
    )
    add_body(
        document,
        "Two design decisions are especially significant. First, the project uses serialized artifacts for the preprocessor and models, ensuring that training-time transformations and inference-time transformations are consistent. This directly supports the thesis objective of reproducibility and trustworthy evaluation. Second, the prediction service adopts a Singleton pattern for model access in web mode through get_instance. This reduces repeated disk I/O and deserialization overhead when handling multiple requests, while preserving model-switch support by re-instantiating when a different model slug is requested."
    )
    add_body(
        document,
        "The interaction pattern can be described as request-driven deterministic flow. The frontend obtains the form schema and model list from dedicated endpoints, builds the UI dynamically, and submits one structured applicant record. The backend then executes an invariant sequence: validate input schema and ranges, transform by fitted pipeline, compute class and probability, derive risk level, then compute explanation vectors. This sequencing enforces a stable contract between modules and avoids route-level code duplication. The complete class and module relationships are summarised in Figure 3.1."
    )
    add_captioned_image(
        document,
        FIGURE_3_1_PATH,
        "Thesis Class UML Diagram",
        "UML class diagram derived from the implemented codebase. It shows the configuration objects, Flask application layer, CreditRiskPredictor orchestrator, preprocessing and training modules, validation logic, explainability modules, and the wrapper used for LIME explanations."
    )

    add_heading(document, "3.2 Technology Stack", size=13)
    add_body(
        document,
        "Python is selected as the implementation language because it offers mature libraries for machine learning, interpretability, and scientific data handling in a single ecosystem. Flask is chosen as the backend framework due to its lightweight routing model and low ceremony, which is appropriate for an academic prototype where control over endpoint behavior is preferred over framework abstraction. The frontend is intentionally minimal (HTML, CSS, JavaScript with Chart.js) to keep the focus on model interaction quality rather than framework complexity."
    )
    add_body(
        document,
        "For modelling, scikit-learn is used for preprocessing, classic estimators, and cross-validation utilities, while XGBoost is included for high-performing boosted trees on tabular data. SHAP and LIME are both integrated for explainability, but with different roles: SHAP is in the primary prediction response path, whereas LIME serves as a complementary local explanation method and experimentation utility. Pandas and NumPy provide data-frame and array-level operations that make the pipeline testable and deterministic."
    )
    add_body(
        document,
        "This stack reflects a deliberate trade-off aligned with Chapter 1 objectives: maximize interpretability and reproducibility without introducing unnecessary platform complexity such as a database server or frontend SPA framework. Because the dataset is static and bounded, file-based artifact persistence through NumPy and joblib is sufficient and keeps deployment simple for examination environments."
    )

    add_heading(document, "3.3 Data Model and Structures", size=13)
    add_body(
        document,
        "The application does not use a relational database schema in its current form. Instead, it uses a file-backed data model and in-memory tabular structures. The raw source is the German Credit file with 20 input attributes and one target attribute. Inputs are partitioned into two explicit lists: numerical columns and categorical columns. This separation is not cosmetic; it drives the ColumnTransformer branches and ensures type-aware preprocessing. Category mapping dictionaries convert encoded domain codes into human-readable labels, which are then reused by validation and schema generation to guarantee frontend-backend consistency."
    )
    add_body(
        document,
        "Primary structures include: DataFrame rows for raw and validated records; NumPy arrays for transformed train/test matrices; serialized preprocessor and model objects as joblib artifacts; and structured JSON dictionaries for API contracts. Validation constraints are represented as dictionary-based range and option sets, yielding explicit and inspectable rule definitions. The schema endpoint exports these structures into frontend-consumable field descriptors, which is effectively a lightweight metadata model for dynamic form generation."
    )
    add_body(
        document,
        "Key classes and objects in the developer data model are the CreditRiskPredictor service object, the fitted ColumnTransformer pipeline, estimator instances for Logistic Regression, Random Forest, and XGBoost, and explainer objects generated per model family. Together they form a typed transformation graph from user payload to interpreted decision output. The transformation pipeline is summarised in Figure 3.2."
    )
    add_image_placeholder(
        document,
        "Data Transformation Pipeline Diagram",
        "Provide a diagram showing raw JSON input, validation, preprocessing branches, the merged feature matrix, model inference, and SHAP explanation output."
    )

    add_heading(document, "3.4 Implementation Logic and Key Algorithms", size=13)
    add_body(
        document,
        "The computational engine of the application is the predict method of the CreditRiskPredictor module. This method operationalizes a multi-stage decision function where every stage has explicit failure behavior. Stage 1 performs schema and semantic validation, preventing malformed requests from reaching model inference. Stage 2 applies the persisted preprocessing pipeline to enforce training-inference parity. Stage 3 performs probabilistic inference via predict_proba and class inference via predict. Stage 4 computes explanation vectors and extracts top positive and negative feature contributors. The final output object merges machine output and interpretability metadata into one stable response contract."
    )
    add_body(
        document,
        "The training engine is similarly modular. Data is split with stratification to preserve class distribution in train and test sets, avoiding optimistic evaluation due to imbalance distortion. The preprocessing pipeline uses imputation and encoding before model fit, and optional SMOTE can be applied only to the training subset to avoid leakage. Model training supports optional hyperparameter tuning for tree-based models using GridSearchCV. Evaluation computes both threshold-based and ranking-based metrics, then supplements holdout metrics with stratified cross-validation summary statistics to improve robustness of model comparison."
    )
    add_body(
        document,
        "Explainability logic includes model-sensitive explainer selection. Tree models use TreeExplainer when possible; logistic regression uses LinearExplainer; and compatibility fallback can use KernelExplainer with background summarization. This strategy acknowledges that interpretability algorithms have model-dependent assumptions and runtime profiles. Therefore, explanation is treated as a first-class algorithmic branch, not as a visualization afterthought."
    )

    add_pseudocode_block(
        document,
        "Pseudocode: End-to-End Prediction Pipeline",
        [
            "FUNCTION Predict(raw_input, model_name, explain=True):",
            "    model_service <- CreditRiskPredictor.get_instance(model_name)",
            "    validated_df, errors <- validate_input(raw_input)",
            "    IF errors NOT EMPTY: RETURN error_response(errors)",
            "    x_transformed <- model_service.preprocessor.transform(validated_df)",
            "    probability <- model.predict_proba(x_transformed)[class_1]",
            "    prediction <- model.predict(x_transformed)",
            "    risk <- map_probability_to_risk_level(probability)",
            "    response <- {prediction, probability, risk, model_name}",
            "    IF explain:",
            "        explanation <- shap_explain(model, x_transformed[0], feature_names, background)",
            "        response.explanation <- top_positive_negative(explanation)",
            "    RETURN success_response(response)",
        ],
    )

    add_pseudocode_block(
        document,
        "Pseudocode: Training and Evaluation Workflow",
        [
            "FUNCTION TrainAndEvaluate(use_smote=False, tune=False):",
            "    X_train, X_test, y_train, y_test, preprocessor <- prepare_data(use_smote)",
            "    models <- {LogisticRegression, RandomForest, XGBoost(scale_pos_weight)}",
            "    FOR each model IN models:",
            "        IF tune AND model supports grid-search: model <- GridSearchCV(model)",
            "        fit(model, X_train, y_train)",
            "        metrics_test <- evaluate_model(model, X_test, y_test)",
            "        metrics_cv <- cross_validate_model(fresh_model_instance, X_all, y_all)",
            "        persist(model)",
            "    RETURN ranked_models_by_metric(metrics_test, metric='roc_auc')",
        ],
    )

    add_heading(document, "3.5 API and Module Documentation", size=13)
    add_body(
        document,
        "The public API layer consists of five Flask routes. The root route serves the interface template. The schema route returns input metadata for all required fields, including numerical bounds and categorical options. The models route lists persisted model slugs available for inference. The predict route accepts JSON input, optionally with model selection, and returns status, decision outputs, and explanation payload. The health route returns a minimal heartbeat status for operational checks."
    )
    add_body(
        document,
        "Module-level responsibilities are intentionally strict. The data_loader module encapsulates dataset column definitions and categorical decoding. The preprocessor module encapsulates transformation pipeline creation, splitting strategy, and artifact persistence. The train_models module defines model factories, training orchestration, tuning branch logic, and model loading utilities. The evaluate module provides metrics and cross-validation abstractions. The prediction subsystem encapsulates validation and orchestration logic, while explainability modules isolate SHAP and LIME mechanics. This cohesion reduces side effects and supports targeted unit testing."
    )
    add_body(
        document,
        "The frontend JavaScript layer consumes backend contracts instead of duplicating domain rules. It loads model and schema metadata, serializes form values to JSON, sends asynchronous prediction requests, and renders response data into decision banners and chart views. This contract-driven frontend design minimizes mismatch risk and simplifies future schema evolution."
    )

    add_heading(document, "3.6 Security and Optimization", size=13)
    add_body(
        document,
        "From a security perspective, the dominant controls in the current release are input validation, constrained schema options, and controlled error handling. Validation ensures required fields, type correctness, and bounded numerical ranges. Categorical values are checked against explicit whitelists, reducing the chance of malformed payload propagation. The API returns structured error objects rather than failing silently, improving diagnosability while preserving endpoint stability."
    )
    add_body(
        document,
        "The system does not yet implement authentication, authorization, CSRF protection, or encrypted secret management beyond configurable environment variables. For a production extension, these should be prioritized together with request rate limiting and audit logging. Nevertheless, for thesis scope and local deployment, the present controls are adequate to demonstrate integrity at the data-validation and pipeline-consistency levels."
    )
    add_body(
        document,
        "Performance optimization decisions are explicit in the architecture: singleton-based model reuse avoids repeated model deserialization, preprocessing artifacts are persisted and reused, and SHAP background sampling caps explainability cost. The training subsystem supports parallelism through estimator-level n_jobs and GridSearchCV parallel execution. These optimizations target the real bottlenecks in tabular ML applications: repeated initialization overhead and expensive explanation computations."
    )
    add_body(
        document,
        "In relation to Chapter 1 goals, the architecture achieves a balanced engineering profile: it is sufficiently rigorous to support reproducible scientific discussion, yet sufficiently lightweight for demonstrable user-facing deployment. The resulting developer design is coherent around one central principle: transparent, testable, and explainable credit-risk decision support."
    )


def main():
    document = Document(DOC_PATH)
    append_chapter_3(document)
    try:
        document.save(DOC_PATH)
        print(DOC_PATH)
    except PermissionError:
        target = Path(DOC_PATH)
        fallback = target.with_name("ELTE_Thesis_Documentation_Updated.docx")
        document.save(str(fallback))
        print(str(fallback))


if __name__ == "__main__":
    main()