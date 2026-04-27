from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


OUTPUT_PATH = r"c:\Users\kafis\OneDrive - Eotvos Lorand Tudomanyegyetem\Asztal\Kafi Thesis\credit-risk-prediction-ml\documentation\ELTE_Thesis_Documentation.docx"


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)

    normal_element = normal.element
    rpr = normal_element.get_or_add_rPr()
    rfonts = rpr.rFonts if rpr.rFonts is not None else OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    rpr.append(rfonts) if rfonts.getparent() is None else None


def format_paragraph(paragraph, *, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, space_after=6):
    paragraph.alignment = alignment
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(space_after)
    for run in paragraph.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold


def add_heading(document: Document, text: str, level_size: int = 14):
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(level_size)
    return paragraph


def add_body_paragraph(document: Document, text: str):
    paragraph = document.add_paragraph()
    paragraph.add_run(text)
    format_paragraph(paragraph)
    return paragraph


def build_document() -> Document:
    document = Document()
    set_document_defaults(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.line_spacing = 1.5
    title.paragraph_format.space_after = Pt(12)
    run = title.add_run("Chapter 1: Introduction and Motivation")
    run.bold = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(14)

    add_heading(document, "1.1 Introduction")
    add_body_paragraph(
        document,
        "Credit risk assessment is one of the most consequential analytical tasks in modern financial services, because the decision to approve, reject, or further investigate a credit application directly affects both institutional profitability and individual access to financial opportunity. From a broader perspective, the problem belongs to the wider field of decision support under uncertainty, where predictive systems are expected not only to estimate the likelihood of an adverse outcome, but also to do so in a way that is operationally efficient and societally defensible. In banking, the stakes are particularly high: an inaccurate approval decision can lead to financial loss and portfolio deterioration, while an overly conservative rejection policy can exclude qualified applicants and reduce the fairness of lending practice."
    )
    add_body_paragraph(
        document,
        "Historically, credit evaluation was dominated by rule-based scorecards and expert judgment. These approaches remain valuable because they are transparent, auditable, and aligned with regulatory expectations. However, they are often limited in their ability to capture complex, non-linear interactions among applicant characteristics. Modern machine learning methods offer a stronger predictive capacity by learning patterns directly from data, but this improvement typically comes at the cost of interpretability. The present project is situated exactly within this tension: it aims to build a system that combines predictive modelling with explanation mechanisms so that the model output can be understood, questioned, and communicated in practical lending contexts."
    )
    add_body_paragraph(
        document,
        "The implemented software addresses this challenge through a complete pipeline built around the German Credit dataset. The dataset is suitable for academic and methodological study because it contains a mixture of numerical and categorical attributes that reflect typical loan-application information. In the codebase, the data is first normalized into a human-readable schema, then transformed by a preprocessing pipeline, and finally used to train and evaluate several supervised learning models. This structure is important not only from an engineering standpoint, but also from a research perspective: it allows the thesis to discuss how data preparation, model choice, and explanation strategy jointly shape the quality of a credit-risk system."
    )

    add_heading(document, "1.2 Problem Statement")
    add_body_paragraph(
        document,
        "The central problem targeted by the project is the gap between predictive performance and decision transparency in credit-risk modelling. Many automated scoring systems are optimized primarily for accuracy or ranking ability, yet a lender cannot rely on numerical performance alone when the decision affects financial responsibility and customer trust. A model that produces strong metrics on historical data may still be hard to use in practice if it cannot explain why a given applicant is classified as high risk, which attributes influenced the prediction, or whether the result changes under alternative modelling assumptions."
    )
    add_body_paragraph(
        document,
        "This gap is visible in the code architecture. The system does not stop at model training; it also includes an input validator, a schema generator for dynamic forms, and SHAP- and LIME-based explanation layers. These components indicate that the project is not merely a classifier, but a decision-support application that has to accept structured user input, enforce schema consistency, and return interpretable output. Without such mechanisms, the system would be vulnerable to common practical failures: malformed user input, inconsistent preprocessing between training and inference, and outputs that are mathematically valid but unusable for explanation-driven human decision making."
    )
    add_body_paragraph(
        document,
        "A further contradiction in existing solutions is that many business-facing machine learning applications treat interpretability as an optional add-on rather than a design requirement. In the present project, explainability is integrated into the inference path itself. The web endpoint returns not only a class prediction and probability score, but also a risk band and feature-level contributions derived from SHAP. This design choice reflects a deeper problem in applied machine learning: if the system cannot communicate the reasoning behind a result, then the decision may be technically correct yet operationally weak, because users cannot validate, challenge, or trust it effectively."
    )

    add_heading(document, "1.3 Motivation")
    add_body_paragraph(
        document,
        "The motivation for building this software is both methodological and practical. From a methodological standpoint, credit-risk prediction is a canonical classification task that allows the thesis to examine the full machine-learning lifecycle: data acquisition, feature engineering, model comparison, evaluation, and post-hoc explanation. From a practical standpoint, lending institutions increasingly need tools that support rapid decisions without sacrificing accountability. This is especially relevant in environments where an analyst, loan officer, or customer service representative must justify a decision in real time and must be able to identify the dominant factors behind the system output."
    )
    add_body_paragraph(
        document,
        "The codebase shows that the project was motivated by the need to move beyond a single-model demonstration. Three model families are implemented: logistic regression, random forest, and XGBoost. This is not a superficial comparison. Each model represents a different point in the trade-off space between interpretability, non-linearity, and ensemble performance. Logistic regression offers a comparatively transparent linear baseline; random forest captures feature interactions while retaining tree-based explainability tools; XGBoost provides a more optimized gradient-boosted alternative that is often stronger on tabular data. By implementing all three, the project can compare not only predictive scores but also the practical implications of different modelling philosophies."
    )
    add_body_paragraph(
        document,
        "Another source of motivation is the educational value of making the complete system inspectable. The repository separates preprocessing, training, prediction, explainability, and web presentation into distinct modules. This architecture supports a thesis narrative in which every design decision can be justified: why the target is encoded as good versus bad credit, why categorical variables are decoded into readable labels, why the preprocessing pipeline is fitted only on training data, and why SHAP explanations are computed at inference time with a background sample. For an academic project at ELTE, this kind of traceability is especially important because it demonstrates not only that the software works, but that its behaviour is grounded in reproducible engineering choices."
    )

    add_heading(document, "1.4 Background")
    add_body_paragraph(
        document,
        "The technical background of the project combines classical supervised learning, tabular data preprocessing, and explainable artificial intelligence. The problem is framed as binary classification, where the target variable distinguishes between acceptable and risky credit outcomes. This setting is well matched to scikit-learn, whose pipeline abstractions make it possible to encode the full preprocessing logic as a reusable, serializable transformation. In the code, numerical variables are imputed with medians and standardized, while categorical variables are imputed with a constant placeholder and one-hot encoded. This combination is standard for mixed-type tabular data because it reduces the effect of missing values, preserves numerical scale for linear models, and converts nominal categories into a machine-readable representation."
    )
    add_body_paragraph(
        document,
        "The selected dataset, the German Credit dataset from the UCI repository, is widely used in academic credit-scoring literature because it is compact, well documented, and sufficiently realistic to support methodological comparison. Its structure is valuable for studying the relationship between applicant attributes and default risk, but it also introduces typical engineering issues: class imbalance, heterogeneous feature types, and the need to preserve the semantics of domain-specific categories. The preprocessing layer in the project responds to these issues by providing a consistent schema, a stratified train-test split, optional SMOTE-based balancing for training, and serialization of the fitted preprocessor so that the same transformation can be applied during inference."
    )
    add_body_paragraph(
        document,
        "The explainability background is equally important. SHAP is used because it is grounded in cooperative game theory and provides feature-attribution values that sum to the difference between the model output and a baseline expectation. This is suitable for both global and local explanation, which makes it particularly useful in a credit-risk context. The codebase supports TreeExplainer for tree-based models, LinearExplainer for logistic regression, and a KernelExplainer fallback when compatibility issues arise. LIME is also included as a complementary technique because it approximates a local decision boundary with an interpretable surrogate model. The presence of both methods reflects an important theoretical point: explainability in practice is rarely achieved through one universally best algorithm; instead, the choice of explainer must follow the model family, the computational budget, and the intended audience."
    )
    add_body_paragraph(
        document,
        "The project’s web stack uses Flask, HTML, CSS, and JavaScript to expose the model as an interactive application. The frontend does not hard-code the form schema; it queries the backend for the current list of models and the structured input specification. This is a sound design pattern because it reduces duplication and helps the UI stay synchronized with the backend schema. The application therefore demonstrates an important applied software-engineering principle: predictive systems are not only models, but also interfaces, validation layers, storage conventions, and explanation services that must remain consistent over time."
    )

    add_heading(document, "1.5 Aims and Objectives")
    add_body_paragraph(
        document,
        "The first objective is to develop an end-to-end credit-risk prediction pipeline that can reliably transform raw applicant data into a machine-learning prediction. This includes loading the dataset, decoding the categorical codes into readable labels, constructing a stable preprocessing pipeline, and persisting both processed data and trained models for later reuse."
    )
    add_body_paragraph(
        document,
        "The second objective is to compare multiple supervised learning approaches on the same problem setting and to evaluate them using metrics that are meaningful for imbalanced classification. The implementation of logistic regression, random forest, and XGBoost makes it possible to discuss the relative strengths of linear and ensemble methods on structured financial data, not only in terms of accuracy, but also in terms of recall, precision, F1 score, ROC-AUC, and precision-recall behaviour."
    )
    add_body_paragraph(
        document,
        "The third objective is to ensure that predictions are interpretable and operationally usable. To achieve this, the software generates SHAP-based feature contributions for individual cases and presents the result through a web interface that can be used without direct interaction with the Python backend. The system therefore aims to support both model analysis and decision communication, which is essential for a thesis project focused on explainable AI in finance."
    )

    add_heading(document, "1.6 Structure of the Documentation")
    add_body_paragraph(
        document,
        "The remainder of this thesis documentation is organised into three chapters. Chapter 2 provides user documentation, explaining how the application is installed, how the data preparation and model workflow are executed, and how the web interface is used in practice. Chapter 3 presents developer documentation, including the software architecture, the preprocessing and training design, the explanation mechanisms, and the implementation details that make the system reproducible. Chapter 4 concludes the thesis by interpreting the results, reflecting on the strengths and limitations of the approach, and summarising the main findings together with possible directions for future work."
    )

    return document


def main() -> None:
    document = build_document()
    document.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()