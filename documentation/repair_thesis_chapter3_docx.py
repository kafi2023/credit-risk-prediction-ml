from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


DOC_PATH = Path(__file__).with_name("ELTE_Thesis_Documentation_Updated.docx")
PRIMARY_DOC_PATH = Path(__file__).with_name("ELTE_Thesis_Documentation.docx")
FIXED_DOC_PATH = Path(__file__).with_name("ELTE_Thesis_Documentation_Chapter3Fixed.docx")
CLASS_UML_PATH = Path(__file__).with_name("thesis_class_uml.png")
PIPELINE_PATH = Path(__file__).with_name("data_pipeline_diagram.png")


def _set_run_font(run, *, bold: bool = False, size: int = 12) -> None:
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)


def _delete_paragraph(paragraph) -> None:
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)


def _find_paragraph_index(document: Document, text: str, *, occurrence: int = 1) -> int:
    matches = [index for index, paragraph in enumerate(document.paragraphs) if paragraph.text.strip() == text]
    if len(matches) < occurrence:
        raise ValueError(f"Could not find paragraph: {text!r} occurrence={occurrence}")
    return matches[occurrence - 1]


def _add_captioned_image(document: Document, paragraph_index: int, image_path: Path, figure_number: str, title: str, caption: str) -> None:
    paragraph = document.paragraphs[paragraph_index]
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    picture_run = paragraph.add_run()
    picture_run.add_picture(str(image_path), width=Inches(6.3))

    caption_paragraph = document.paragraphs[paragraph_index + 1]
    caption_paragraph.text = ""
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_paragraph.paragraph_format.line_spacing = 1.0
    caption_paragraph.paragraph_format.space_after = Pt(6)

    caption_run = caption_paragraph.add_run(f"Figure {figure_number}: {title}\n")
    _set_run_font(caption_run, bold=True, size=11)
    body_run = caption_paragraph.add_run(caption)
    _set_run_font(body_run, bold=False, size=10)


def main() -> None:
    if not DOC_PATH.exists():
        raise FileNotFoundError(DOC_PATH)

    document = Document(str(DOC_PATH))

    first_chapter3 = _find_paragraph_index(document, "Chapter 3: Developer Documentation", occurrence=1)
    second_chapter3 = _find_paragraph_index(document, "Chapter 3: Developer Documentation", occurrence=2)

    for index in range(second_chapter3 - 1, first_chapter3 - 1, -1):
        _delete_paragraph(document.paragraphs[index])

    figure3_1_text = "The complete class and module relationships are summarised in Figure 3.1."
    for paragraph in document.paragraphs:
        if paragraph.text.strip() == figure3_1_text:
            paragraph.text = figure3_1_text
            break

    figure3_2_reference = "The data transformation pipeline is summarised in Figure 3.2."
    for paragraph in document.paragraphs:
        if paragraph.text.strip().startswith("Key classes and objects in the developer data model are"):
            paragraph.text = paragraph.text.rstrip(".") + ". " + figure3_2_reference
            break

    placeholder_text = "[IMAGE PLACEHOLDER: ER Diagram or Database Schema]"
    placeholder_index = _find_paragraph_index(document, placeholder_text, occurrence=1)
    _add_captioned_image(
        document,
        placeholder_index,
        PIPELINE_PATH,
        "3.2",
        "Data Transformation Pipeline",
        "End-to-end data flow from raw JSON input to validated tabular data, parallel preprocessing branches, merged feature matrix, model inference, and SHAP-based explanation output.",
    )

    for target in (FIXED_DOC_PATH, DOC_PATH, PRIMARY_DOC_PATH):
        try:
            document.save(str(target))
        except PermissionError:
            continue

    print(str(FIXED_DOC_PATH))


if __name__ == "__main__":
    main()