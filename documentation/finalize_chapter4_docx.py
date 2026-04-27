import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


DOC_DIR = Path(r"c:\Users\kafis\OneDrive - Eotvos Lorand Tudomanyegyetem\Asztal\Kafi Thesis\credit-risk-prediction-ml\documentation")
PRIMARY_DOC = DOC_DIR / "ELTE_Thesis_Documentation.docx"
UPDATED_DOC = DOC_DIR / "ELTE_Thesis_Documentation_Updated.docx"
FINAL_DOC = DOC_DIR / "ELTE_Thesis_Documentation_Final.docx"


def _set_run_font(run, *, bold=False, size=12):
    run.bold = bold
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    r = run._element
    rpr = r.get_or_add_rPr()
    rfonts = rpr.rFonts if rpr.rFonts is not None else OxmlElement("w:rFonts")
    rfonts.set(qn("w:ascii"), "Times New Roman")
    rfonts.set(qn("w:hAnsi"), "Times New Roman")
    rfonts.set(qn("w:cs"), "Times New Roman")
    if rfonts.getparent() is None:
        rpr.append(rfonts)


def _format_paragraph(paragraph, *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, after=6):
    paragraph.alignment = align
    paragraph.paragraph_format.line_spacing = line_spacing
    paragraph.paragraph_format.space_after = Pt(after)


def add_heading(document, text, size=14):
    p = document.add_paragraph()
    _format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, after=6)
    run = p.add_run(text)
    _set_run_font(run, bold=True, size=size)


def add_body(document, text):
    p = document.add_paragraph()
    _format_paragraph(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5, after=6)
    run = p.add_run(text)
    _set_run_font(run, bold=False, size=12)


def choose_source_document() -> Path:
    if UPDATED_DOC.exists():
        return UPDATED_DOC
    return PRIMARY_DOC


def apply_heading_styles_for_toc(document: Document) -> None:
    for p in document.paragraphs:
        text = (p.text or "").strip()
        if not text:
            continue

        if re.match(r"^Chapter\s+\d+\s*:", text):
            p.style = document.styles["Heading 1"]
        elif re.match(r"^\d+\.\d+\.\d+\s+", text):
            p.style = document.styles["Heading 3"]
        elif re.match(r"^\d+\.\d+\s+", text):
            p.style = document.styles["Heading 2"]
        else:
            continue

        _format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, after=6)
        for run in p.runs:
            _set_run_font(run, bold=True, size=12)


def insert_toc_at_start(document: Document) -> None:
    title = document.add_paragraph()
    _format_paragraph(title, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, after=6)
    title_run = title.add_run("Table of Contents")
    _set_run_font(title_run, bold=True, size=14)

    toc_paragraph = document.add_paragraph()
    _format_paragraph(toc_paragraph, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, after=6)

    run = toc_paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    hint = OxmlElement("w:t")
    hint.text = "Right-click this field and choose 'Update Field' to refresh the full table of contents."

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_sep)
    r_element.append(hint)
    r_element.append(fld_end)

    spacer = document.add_paragraph()
    _format_paragraph(spacer, align=WD_ALIGN_PARAGRAPH.LEFT, line_spacing=1.5, after=6)

    body = document._body._element
    body.insert(0, spacer._p)
    body.insert(0, toc_paragraph._p)
    body.insert(0, title._p)


def append_chapter_4(document: Document) -> None:
    add_heading(document, "Chapter 4: Conclusion", size=14)

    add_heading(document, "4.1 Summary of Results", size=13)
    add_body(
        document,
        "The project achieved the three core objectives set in Chapter 1 through an integrated technical implementation. First, it delivered a complete end-to-end prediction system: user input is validated, transformed by the persisted preprocessing pipeline, evaluated by a selectable trained model, and returned with an interpretable risk output. Second, it implemented comparative modelling with Logistic Regression, Random Forest, and XGBoost, together with holdout and cross-validated evaluation metrics, enabling methodologically meaningful model selection rather than single-model reporting. Third, it operationalized explainable AI in the runtime path by attaching SHAP-based feature contributions to each prediction response, thereby addressing the transparency requirement that motivated the thesis."
    )
    add_body(
        document,
        "At software-engineering level, these results are supported by modularization and reproducibility controls. Data loading, preprocessing, training, evaluation, prediction, and explainability were separated into dedicated modules with focused responsibilities. Serialized artifacts (preprocessor and model objects) ensured consistency between training and inference phases, and the web interface consumed backend schema metadata dynamically, reducing form drift and maintaining contract alignment. This confirms that the practical software objective was not only algorithmic performance, but also deployable and inspectable behavior for non-technical stakeholders."
    )

    add_heading(document, "4.2 Critical Evaluation and Discussion", size=13)
    add_body(
        document,
        "The most significant achievement is the synthesis of predictive modelling and interpretability in one coherent application workflow. Many baseline credit-scoring demonstrations terminate at probability output, whereas this system delivers a user-facing explanation artifact for each prediction. This is a material improvement over opaque scoring prototypes because decision consumers can inspect directional feature impact instead of trusting a scalar score in isolation. The inclusion of model alternatives further strengthens the contribution by exposing trade-offs among linear transparency, ensemble robustness, and boosted-tree performance."
    )
    add_body(
        document,
        "Compared with the background solutions discussed in Chapter 1, the project occupies an intermediate but valuable position. It is more transparent than black-box-only implementations, yet more expressive than rigid rule-based scoring approaches that cannot model higher-order interactions effectively. From a research perspective, this balance is important: the work demonstrates that explainability does not need to be postponed to offline analytics, but can be integrated into the operational inference path. From a software perspective, the architecture shows that this integration is feasible with modest infrastructure using Python and Flask, which is relevant for academic and small-team deployment contexts."
    )

    add_heading(document, "4.3 Limitations", size=13)
    add_body(
        document,
        "Several limitations remain and should be interpreted carefully. The dataset is a standard benchmark with fixed size and historical context, which constrains external validity when extrapolating to contemporary multi-source banking environments. The current application also lacks authentication and role-based authorization, so all users with endpoint access share identical operational privileges. For thesis scope this was acceptable, but for institutional deployment it is insufficient."
    )
    add_body(
        document,
        "From a technical standpoint, explanation latency can become a bottleneck when SHAP computations are requested at high throughput, especially for more computationally expensive explainers or larger background sets. Model lifecycle governance is also minimal: there is no automated drift detection, no model version registry, and no persistent prediction audit trail. Error handling is structured and user-visible, but observability remains lightweight without centralized logging, tracing, or alerting. Finally, frontend interaction is intentionally simple and does not yet offer workflow features such as batch inference, saved scenarios, or report export for audit use."
    )

    add_heading(document, "4.4 Future Work", size=13)
    add_body(
        document,
        "Future development can proceed along four dimensions. First, governance and security: add authentication, role-aware permissions, request-level audit logging, and stricter secret/key management. Second, MLOps maturity: introduce model versioning, automated retraining triggers, drift monitoring, and reproducible experiment tracking. Third, scalability and architecture: migrate from single-process Flask serving to containerized deployment with reverse proxying, horizontal scaling, and asynchronous task offloading for expensive explanation jobs. Fourth, user-facing capability: support batch uploads, downloadable explanation reports, and configurable threshold policies aligned with institutional risk appetite."
    )
    add_body(
        document,
        "Methodologically, future research can also evaluate calibration quality, fairness-sensitive metrics across protected groups, and scenario-based stress testing under distribution shift. Alternative modelling approaches such as calibrated gradient boosting variants, monotonic constraints, or hybrid scorecard-ML ensembles may further improve the balance between interpretability and predictive stability. These extensions would preserve the thesis core principle while extending practical relevance for production credit ecosystems."
    )

    add_heading(document, "4.5 Concluding Remarks", size=13)
    add_body(
        document,
        "This thesis contributes to the identified niche by demonstrating a concrete, working bridge between accurate tabular credit-risk prediction and explainable decision support. The final system is not merely a model benchmark; it is an integrated software artifact that operationalizes validation, inference, explanation, and human-readable interaction in one pipeline. The resulting contribution is therefore both technical and methodological: it offers a reproducible implementation pattern for explainable credit analytics and provides a foundation for subsequent academic and industrial extensions."
    )


def append_reference_list(document: Document) -> None:
    add_heading(document, "Reference List", size=14)
    references = [
        "Biecek, P., & Burzykowski, T. (2021). Explanatory Model Analysis: Explore, Explain, and Examine Predictive Models. Chapman and Hall/CRC.",
        "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5-32.",
        "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting system. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 785-794.",
        "Dua, D., & Graff, C. (2019). UCI Machine Learning Repository [German Credit Data]. University of California, Irvine, School of Information and Computer Sciences. http://archive.ics.uci.edu/ml",
        "Flask Documentation. (2026). Flask: Web development, one drop at a time. https://flask.palletsprojects.com/",
        "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to interpreting model predictions. Advances in Neural Information Processing Systems, 30.",
        "Molnar, C. (2022). Interpretable Machine Learning (2nd ed.). https://christophm.github.io/interpretable-ml-book/",
        "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., Grisel, O., Blondel, M., Prettenhofer, P., Weiss, R., Dubourg, V., Vanderplas, J., Passos, A., Cournapeau, D., Brucher, M., Perrot, M., & Duchesnay, E. (2011). Scikit-learn: Machine learning in Python. Journal of Machine Learning Research, 12, 2825-2830.",
        "Ribeiro, M. T., Singh, S., & Guestrin, C. (2016). Why should I trust you? Explaining the predictions of any classifier. Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge Discovery and Data Mining, 1135-1144.",
        "Wes McKinney. (2010). Data structures for statistical computing in Python. Proceedings of the 9th Python in Science Conference, 56-61.",
    ]
    for ref in references:
        add_body(document, ref)


def append_ai_appendix(document: Document) -> None:
    add_heading(document, "Appendix: AI Usage Declaration", size=14)
    add_body(
        document,
        "This appendix declares the use of AI-assisted drafting support during thesis documentation preparation. AI support was used to organize chapter structure, produce candidate wording, and transform codebase observations into technical prose. The final scientific interpretation, argument selection, and acceptance of written content remained under student control."
    )

    add_heading(document, "A.1 Prompt Summary by Chapter", size=13)
    add_body(
        document,
        "Chapter 1 prompts focused on establishing introduction, motivation, problem framing, background, and objective alignment based on real code architecture and thesis requirements. Purpose: structuring and initial drafting of academic narrative from validated implementation details."
    )
    add_body(
        document,
        "Chapter 2 prompts focused on user documentation workflows, UI behavior, endpoint-driven interaction, setup steps, and troubleshooting guidance, including mandatory screenshot placeholders and capture instructions. Purpose: practical procedural writing for non-technical users."
    )
    add_body(
        document,
        "Chapter 3 prompts focused on developer-level synthesis of architecture, module responsibilities, data flow, algorithmic pipeline logic, and design rationale with implementation-aware depth. Purpose: technical synthesis and critical architectural explanation."
    )
    add_body(
        document,
        "Chapter 4 prompts focused on critical reflection, limitations, future work, and academic closure, plus document-wide finalization tasks (TOC, references, and declaration). Purpose: evaluative interpretation and completion of the thesis narrative."
    )

    add_heading(document, "A.2 Verification Statement", size=13)
    add_body(
        document,
        "The student author reviewed and verified the final interpretations, technical claims, and conclusions before submission. AI-generated text was treated as draft assistance, and responsibility for correctness, citation quality, and academic integrity remains with the student."
    )


def main() -> None:
    source = choose_source_document()
    document = Document(str(source))

    append_chapter_4(document)
    append_reference_list(document)
    append_ai_appendix(document)
    apply_heading_styles_for_toc(document)
    insert_toc_at_start(document)

    save_targets = [PRIMARY_DOC, UPDATED_DOC, FINAL_DOC]
    saved_to = None
    for target in save_targets:
        try:
            document.save(str(target))
            saved_to = target
            break
        except PermissionError:
            continue

    if saved_to is None:
        fallback = DOC_DIR / "ELTE_Thesis_Documentation_Final_Copy.docx"
        document.save(str(fallback))
        saved_to = fallback

    print(str(saved_to))


if __name__ == "__main__":
    main()