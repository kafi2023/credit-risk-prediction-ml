import os
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.shared import Pt, Inches, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxcompose.composer import Composer
import fitz

ROOT = Path(__file__).parent
PDF_PATH = ROOT / "Kafi MD Abdullah Hel (N06WMD); Request identifier-1653963517; Request submission date-1052025 10518 PM.pdf"
THESIS_BODY = ROOT / 'documentation' / 'ELTE_Thesis_Documentation.docx'
TITLE_PAGE = ROOT / 'new_title_page.docx'
FINAL_OUT = ROOT / 'documentation' / 'Thesis_KafiMDAbdullahHel_Final.docx'
LOGO_PATH = ROOT / 'documentation' / 'logo.jpg'
DIAGRAMS_DIR = ROOT / 'docs' / 'diagrams'


def set_font(run, size=12, bold=False, italic=False, small_caps=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    rPr = run.font.element.get_or_add_rPr()
    if small_caps:
        smallCaps = OxmlElement('w:smallCaps')
        rPr.append(smallCaps)
    rfonts = OxmlElement('w:rFonts')
    rfonts.set(qn('w:ascii'), 'Times New Roman')
    rfonts.set(qn('w:hAnsi'), 'Times New Roman')
    rfonts.set(qn('w:cs'), 'Times New Roman')
    rPr.append(rfonts)


def create_title_page():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.0)
    section.bottom_margin = Cm(3.0)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)

    table = doc.add_table(rows=1, cols=2)
    table.columns[0].width = Cm(4.0)
    table.columns[1].width = Cm(11.0)

    row = table.rows[0]
    cell_logo = row.cells[0]
    cell_logo.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if LOGO_PATH.exists():
        p = cell_logo.add_paragraph()
        run = p.add_run()
        run.add_picture(str(LOGO_PATH), width=Cm(3.5))

    cell_text = row.cells[1]
    cell_text.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell_text.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.1

    r1 = p.add_run('Eötvös Loránd University\n')
    set_font(r1, size=18, small_caps=True)
    r2 = p.add_run('Faculty of Informatics\n')
    set_font(r2, size=18, small_caps=True)
    r3 = p.add_run('Dept. of Software Technology and Methodology')
    set_font(r3, size=14, small_caps=True)

    p_space1 = doc.add_paragraph()
    p_space1.paragraph_format.space_before = Pt(90)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run('Interactive Machine Learning Model Explanation Web Tool for Credit Risk Prediction')
    set_font(r_title, size=18, bold=True)

    p_space2 = doc.add_paragraph()
    p_space2.paragraph_format.space_before = Pt(60)

    table_bottom = doc.add_table(rows=3, cols=3)
    table_bottom.columns[0].width = Cm(6.5)
    table_bottom.columns[1].width = Cm(2.0)
    table_bottom.columns[2].width = Cm(6.5)

    row0 = table_bottom.rows[0]
    r_sup_label = row0.cells[0].paragraphs[0].add_run('Supervisor:')
    set_font(r_sup_label, size=14, italic=True)
    r_auth_label = row0.cells[2].paragraphs[0].add_run('Author:')
    set_font(r_auth_label, size=14, italic=True)

    row1 = table_bottom.rows[1]
    r_sup_name = row1.cells[0].paragraphs[0].add_run('Md. Easin Arafat')
    set_font(r_sup_name, size=14)
    r_auth_name = row1.cells[2].paragraphs[0].add_run('Kafi MD Abdullah Hel')
    set_font(r_auth_name, size=14)

    row2 = table_bottom.rows[2]
    r_sup_aff = row2.cells[0].paragraphs[0].add_run('PhD Candidate')
    set_font(r_sup_aff, size=12)
    r_auth_deg = row2.cells[2].paragraphs[0].add_run('Computer Science BSc')
    set_font(r_auth_deg, size=12)

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.space_before = Pt(60)
    r_footer = p_footer.add_run('Budapest, 2025')
    set_font(r_footer, size=12, italic=True)

    doc.save(str(TITLE_PAGE))


def create_declaration_doc():
    doc = Document()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run('Declaration of Authorship')
    r_title.bold = True
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(16)

    p_text = doc.add_paragraph()
    p_text.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r_text = p_text.add_run(
        "I, Kafi MD Abdullah Hel, computer science BSc student, author of this thesis, declare that except where noted, all work is my own and no part of this thesis has been submitted for a degree to any other university or institution. "
        "Furthermore, I declare that in the preparation of this thesis I have complied with the requirements of the university regarding the avoidance of plagiarism."
    )
    r_text.font.name = 'Times New Roman'
    r_text.font.size = Pt(12)

    p_sign = doc.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_sign = p_sign.add_run("\n\n............................................\nSignature\nBudapest, 2025.")
    r_sign.font.name = 'Times New Roman'
    r_sign.font.size = Pt(12)
    doc.save(str(ROOT / 'temp_declaration.docx'))


def create_pdf_page_doc():
    doc = Document()
    if PDF_PATH.exists():
        pdf = fitz.open(str(PDF_PATH))
        page = pdf[0]
        pix = page.get_pixmap(dpi=150)
        img_path = ROOT / 'temp_pdf_page.png'
        pix.save(str(img_path))
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img_path), width=Inches(6.0))
        doc.save(str(ROOT / 'temp_pdf.docx'))
        return img_path
    return None


def create_uml_doc():
    doc = Document()
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run('Appendix: UML Diagrams')
    r_title.bold = True
    r_title.font.name = 'Times New Roman'
    r_title.font.size = Pt(16)

    diagrams = [
        'System Architecture Diagram.png',
        'component diagram.png',
        'Class Daigram.png',
        'Sequence Diagram - Prediction Flow.png',
        'User Flow Diagram.png',
        'Data Flow.png',
        'state diagram.png'
    ]
    added = False
    for d in diagrams:
        path = DIAGRAMS_DIR / d
        if path.exists():
            added = True
            p_desc = doc.add_paragraph()
            p_desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_desc.add_run(f'\n{d.split(".")[0]}').bold = True
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_img.add_run().add_picture(str(path), width=Inches(5.5))
            doc.add_page_break()
    if added:
        doc.save(str(ROOT / 'temp_uml.docx'))
        return ROOT / 'temp_uml.docx'
    return None


def combine_all():
    create_title_page()
    create_declaration_doc()
    img_path = create_pdf_page_doc()
    uml_doc = create_uml_doc()

    master = Document(str(TITLE_PAGE))
    composer = Composer(master)

    if (ROOT / 'temp_declaration.docx').exists():
        composer.append(Document(str(ROOT / 'temp_declaration.docx')))
    if (ROOT / 'temp_pdf.docx').exists():
        composer.append(Document(str(ROOT / 'temp_pdf.docx')))
    if THESIS_BODY.exists():
        composer.append(Document(str(THESIS_BODY)))
    if uml_doc and uml_doc.exists():
        composer.append(Document(str(uml_doc)))

    master.save(str(FINAL_OUT))
    print(f"Successfully generated {FINAL_OUT}")

    # cleanup temps
    for f in ['temp_declaration.docx', 'temp_pdf.docx', 'temp_pdf_page.png', 'temp_uml.docx', 'new_title_page.docx']:
        p = ROOT / f
        try:
            if p.exists():
                p.unlink()
        except:
            pass


if __name__ == '__main__':
    combine_all()
